# -*- coding: utf-8 -*-
"""
MES智慧管理系统 · Vue3 工业互联网应用系统 —— 期末大作业设计文档生成脚本
版本：v1（2026-06-24）

用途：用 python-docx 生成《设计文档.docx》。
说明：
  - 文档内容全部基于仓库真实代码（mes/vue3），不杜撰功能。
  - 凡需贴图处均以「【截图位置 N：…】」红色占位段落标注，作者自行截图替换。
  - 保留本脚本以便后续增改；如需改版请复制为 gen_doc_v2.py 再修改。
运行：
  cd mes/vue3/docs/design-doc && python3 gen_doc_v1.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "设计文档.docx"

# ------------------------------------------------------------------ 基础样式

doc = Document()

# 中文默认字体（正文宋体，标题黑体）
def set_base_font(document):
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

set_base_font(doc)


def _set_cn_font(run, font="宋体", size=11, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    _set_cn_font(run, font="黑体", size=18, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return p


def h2(text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    _set_cn_font(run, font="黑体", size=15, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
    return p


def h3(text):
    p = doc.add_heading(level=3)
    run = p.add_run(text)
    _set_cn_font(run, font="黑体", size=13, bold=True, color=RGBColor(0x40, 0x40, 0x40))
    return p


def para(text, size=11, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if align:
        p.alignment = align
    run = p.add_run(text)
    _set_cn_font(run, size=size, bold=bold)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    _set_cn_font(run, size=11)
    return p


def code_block(text):
    """等宽代码块，浅灰底。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    # 浅灰底纹
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F3F5")
    p._p.get_or_add_pPr().append(shd)
    return p


def screenshot(caption):
    """红色截图占位段落。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"【截图位置：{caption}】")
    _set_cn_font(run, size=10.5, bold=True, color=RGBColor(0xC0, 0x00, 0x00))
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        _set_cn_font(run, font="黑体", size=10.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            _set_cn_font(run, size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ================================================================== 封面

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(80)
r = title.add_run("MES智慧管理系统")
_set_cn_font(r, font="黑体", size=34, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("基于 Vue3 的智能制造执行系统")
_set_cn_font(r, font="黑体", size=20, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("—— Vue 开发技术 · 期末大作业设计文档 ——")
_set_cn_font(r, size=13)

for _ in range(2):
    doc.add_paragraph()

info_tbl = doc.add_table(rows=0, cols=2)
info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in [
    ("项目名称", "MES智慧管理系统（制造执行系统）Vue3 前端"),
    ("应用领域", "工业互联网 / 智能制造 / 车间生产管理"),
    ("技术栈", "Vue 3.5 + TypeScript + Vite 8 + Pinia + Vue Router + Element Plus"),
    ("代码规模", "126 个单文件组件 / 107 个 TS 模块 / 约 2.1 万行 / 30 个单测文件 333 用例"),
    ("文档版本", "v1"),
    ("完成日期", "2026 年 6 月"),
]:
    cells = info_tbl.add_row().cells
    rk = cells[0].paragraphs[0].add_run(k)
    _set_cn_font(rk, font="黑体", size=11, bold=True)
    rv = cells[1].paragraphs[0].add_run(v)
    _set_cn_font(rv, size=11)

doc.add_page_break()

# ================================================================== 目录提示
h1("目录")
para("（提示：在 Word 中选中下方占位后，通过「引用 → 目录 → 自动目录」可自动生成带页码的目录；"
     "或将光标置于此处插入域目录后按 F9 更新。本文档已使用标准标题样式，可被 Word 正确识别。）",
     size=10)
for i, t in enumerate([
    "一、项目概述",
    "二、需求分析",
    "三、技术选型",
    "四、系统架构",
    "五、核心功能实现细节",
    "六、创新与深度优化",
    "七、代码质量与工程化",
    "八、评分细则对照",
    "九、部署与运行",
    "十、总结与展望",
]):
    para(t, size=11)
doc.add_page_break()

# ================================================================== 一、项目概述
h1("一、项目概述")

para("MES智慧管理系统（Manufacturing Execution System，制造执行系统）是一套面向离散制造车间的"
     "工业互联网应用系统，目标是把工厂从「计划下达」到「车间执行」再到「数据回采」的全过程数字化，"
     "覆盖系统管理、基础数据、工艺技术、生产计划、库存、数字化看板、数字孪生与工作流八大业务域。"
     "本作业以 Vue3 全面重写前端，后端为既有的 Spring Boot 单体（Shiro 会话鉴权），前后端经由 "
     "Vite 开发代理联调，生产环境由 Nginx 反向代理。")

para("选题属于典型的工业互联网场景：制造执行系统是连接企业上层 ERP 与车间底层设备的中枢，"
     "具有明确的角色（管理员、工艺员、计划员、车间班组长、操作工）、清晰的业务流（订单 → 排产 → "
     "派工 → 执行 → 入库）和真实的数据可视化需求（产能看板、3D 仓储孪生）。它既不简单也不脱离实际，"
     "适合充分展示 Vue3 组合式 API、Pinia、Vue Router、组件化与第三方生态集成的综合能力。")

para("本前端是一个纯前后端分离的单页应用（SPA）。它以单文件组件（SFC）为基本单元，"
     "通过 Vue Router 组织多级嵌套路由与动态路由，通过 Pinia 管理用户态、权限态、应用态与通知态，"
     "并以 Element Plus 作为 UI 组件库，叠加 ECharts、Three.js、bpmn-js 等第三方技术栈实现"
     "数据可视化大屏、3D 数字孪生仓库、BPMN 流程设计器、AI 智能助手等多个亮点功能。")

screenshot("系统登录页 + 主界面（侧边栏菜单 + 顶部多页签 + 内容区）整体一览")

# ================================================================== 二、需求分析
h1("二、需求分析")

h2("2.1 业务背景与选题定位")
para("制造执行系统处于工业软件金字塔的中间层：上承企业资源计划（ERP）下达的生产订单，"
     "下接车间设备与人员的实际执行。本系统聚焦车间执行环节，解决三个核心问题：")
bullet("生产要素数字化：物料、设备、零部件、仓库、加工单元、班组员工等基础数据统一建模、统一维护。")
bullet("工艺与计划数字化：工艺路线、产品 BOM、工序内容编制；生产订单录入、排产派工、甘特排程。")
bullet("过程可视与协同：实时产能看板、3D 数字孪生仓库、轻量 BPMN 审批流、站内通知中心与 AI 助手。")

h2("2.2 用户角色与典型场景")
table(
    ["角色", "关注点", "典型操作场景"],
    [
        ["系统管理员", "账号与权限", "维护用户/角色/菜单/部门/字典，分配菜单权限，管理班组与员工"],
        ["工艺工程师", "工艺与 BOM", "定义工序、编排工艺路线、维护产品 BOM、绑定 BOM-工艺、编制工序内容、工艺查询"],
        ["计划员", "订单与排产", "录入生产订单、选择 BOM 派生需求、下达工单、甘特图排程、提交审批"],
        ["车间班组长", "派工与执行", "员工作业派工、查看待办任务、签收/提交/驳回审批节点"],
        ["仓管员", "出入库", "计划入库确认、手动入库、配套出库（FIFO 记账）、库存明细查询"],
        ["管理层", "经营可视", "查看数据可视化大屏、3D 数字孪生仓库热力分布、接收站内通知"],
    ],
    widths=[1.3, 1.6, 4.0],
)

h2("2.3 功能范围（模块清单）")
para("系统按后端八大业务域划分模块，前端共实现 30+ 个业务页面。下表列出主要模块与其代表功能：")
table(
    ["模块", "代表功能", "页面数"],
    [
        ["系统管理 system", "用户、角色、菜单、部门、字典、班组员工、通知中心", "8+"],
        ["基础数据 basedata", "物料（含图片）、设备、设备编组、零部件、加工单元、仓库、动态主数据", "9"],
        ["工艺技术 technology", "工序、工艺路线（穿梭框）、产品 BOM、BOM-工艺、工序内容、工艺查询", "7"],
        ["生产计划 plan / order", "生产订单录入、待办任务、工单下达、派工、生产甘特图", "5"],
        ["库存 inventory", "计划入库确认、手动入库、配套出库、库存明细查询", "4"],
        ["数字化 digitization", "数据可视化大屏（ECharts）、3D 数字孪生仓库（Three.js）", "2"],
        ["工作流 workflow", "流程分类、流程表单、流程定义、BPMN 模型设计器", "4"],
    ],
    widths=[1.8, 4.2, 0.9],
)

h2("2.4 非功能性需求")
bullet("性能：关键交互响应 < 1 秒；首屏加载 < 3 秒；列表分页、路由懒加载、组件按需引入。")
bullet("安全：登录拦截 + 路由级权限校验；会话基于 Shiro Cookie；401 自动跳登录。")
bullet("可用性：加载状态（骨架屏 / Loading）、表单校验与错误反馈、明暗双主题、多页签导航。")
bullet("可维护性：模块化目录、统一 ESLint + Prettier 代码风格、语义化命名、单元测试覆盖关键逻辑。")

# ================================================================== 三、技术选型
h1("三、技术选型")

para("技术选型紧扣「工业互联网应用 + Vue3 综合能力展示」两条主线，所有依赖均为当前主流且长期维护的版本。"
     "下表给出核心技术栈，随后逐项说明选型理由。")

table(
    ["分类", "选型（实装版本）", "作用"],
    [
        ["核心框架", "Vue 3.5.34（<script setup> + 组合式 API）", "单文件组件、响应式、生命周期钩子"],
        ["开发语言", "TypeScript 6.0（vue-tsc 类型检查）", "静态类型、接口约束、可维护性"],
        ["构建工具", "Vite 8.0", "极速冷启动、HMR、按需打包、产物分包"],
        ["路由", "Vue Router 4（实装 5.1，组合式 API 一致）", "多级嵌套路由、动态路由参数、全局守卫"],
        ["状态管理", "Pinia 3.0 + pinia-plugin-persistedstate 4.7", "模块化 store、状态持久化"],
        ["UI 组件库", "Element Plus 2.14 + @element-plus/icons-vue", "企业级表单/表格/弹窗/消息组件"],
        ["按需引入", "unplugin-auto-import + unplugin-vue-components", "API 与组件全自动 tree-shaking"],
        ["HTTP", "Axios 1.18", "拦截器：表单编码 / Result 解包 / 401 处理"],
        ["数据可视化", "ECharts 6 + vue-echarts 8", "产能大屏、统计图表"],
        ["3D 引擎", "Three.js 0.184", "数字孪生仓库三维渲染"],
        ["流程引擎", "bpmn-js 18 + diagram-js-minimap", "BPMN 流程模型可视化设计"],
        ["工具库", "@vueuse/core、dayjs、nprogress、markdown-it", "组合式工具 / 时间 / 进度条 / Markdown"],
        ["动画", "@vueuse/motion、@formkit/auto-animate", "克制的进场动画与列表过渡"],
        ["样式", "Sass + CSS 变量", "明暗双主题驱动"],
        ["测试", "Vitest 4 + @vue/test-utils", "组合式函数 / 纯逻辑单元测试"],
        ["代码规范", "ESLint 10 + Prettier", "统一代码风格"],
    ],
    widths=[1.3, 3.3, 2.4],
)

h3("选型理由")
bullet("Vue 3.5 + 组合式 API：作业明确要求「组合式 API 生命周期钩子」。组合式 API 以 setup 为中心，"
       "逻辑按关注点聚合（如分页、请求、表单各自成 composable），相比 Options API 更利于复用与类型推导。")
bullet("TypeScript：为 30+ 业务页面与 100+ 接口提供静态类型约束，配合 vue-tsc 在构建期拦截类型错误，"
       "显著提升大型项目的可维护性。")
bullet("Vite 8：基于原生 ESM 的开发服务器冷启动毫秒级，HMR 体验优秀；生产构建支持自定义 manualChunks 分包。")
bullet("Vue Router 4：作业要求项。提供嵌套路由、动态路由参数、导航守卫，是 SPA 权限控制的基础设施。")
bullet("Pinia：Vue 官方推荐的状态管理库，API 简洁、TS 友好、天然支持模块化；配合持久化插件实现刷新不丢登录态。")
bullet("Element Plus：作业要求使用第三方 UI 库。其表单、表格、弹窗、消息组件成熟稳定，"
       "且对组合式 API、暗色主题、tree-shaking 支持良好。")
bullet("ECharts / Three.js / bpmn-js：分别支撑数据可视化大屏、3D 数字孪生、BPMN 设计器三大创新亮点，"
       "对应作业「技术栈扩展」与「功能创新」的加分项。")

# ================================================================== 四、系统架构
h1("四、系统架构")

h2("4.1 总体架构")
para("系统采用前后端分离架构。前端为 Vue3 单页应用，运行于 Vite 开发服务器（端口 4200），"
     "通过 /api 前缀的代理把请求转发到后端 Spring Boot（端口 9090）；后端基于 Apache Shiro "
     "做会话级鉴权，登录后通过 Cookie 维持会话。前端请求统一携带 Cookie（withCredentials），"
     "并以 X-Requested-With 头让后端在未授权时返回 401 JSON 而非 HTML 重定向。")

code_block(
    "浏览器(Vue3 SPA :4200)\n"
    "   │  axios / fetch  携带会话 Cookie\n"
    "   ▼\n"
    "Vite Dev Proxy  /api → http://localhost:9090\n"
    "   ▼\n"
    "Spring Boot 单体(:9090)  Shiro 鉴权 + MyBatis-Plus\n"
    "   ▼\n"
    "MySQL 8 / Redis(或 Ehcache)"
)
screenshot("系统总体架构图（前端 SPA / Vite 代理 / 后端 / 数据库分层）")

h2("4.2 前端分层与目录结构")
para("前端遵循 Vite 模块化目录约定，按「职责」与「业务模块」双维度组织：")
code_block(
    "src/\n"
    "├── api/            按后端模块组织的接口函数 + request.ts(axios 封装)\n"
    "├── assets/styles/  主题变量、Element 覆盖、动画\n"
    "├── components/     通用组件(DataTable / SearchForm / FormDialog / EChart …)\n"
    "├── composables/    组合式函数(useRequest / usePagination / useAiChat …)\n"
    "├── layouts/        AdminLayout(后台壳) / ScreenLayout(大屏壳) + 子组件\n"
    "├── plugins/        全局插件注册\n"
    "├── router/         index.ts(路由表) + guards.ts(全局守卫)\n"
    "├── stores/         Pinia 模块(user / permission / app / notice)\n"
    "├── types/          TypeScript 类型声明\n"
    "├── utils/          纯函数工具(权限收集 / SSE 分帧 / BPMN / 甘特 …)\n"
    "├── views/          按业务模块组织的页面组件\n"
    "└── main.ts         应用入口(挂载 Pinia / Router / 守卫 / 插件 / 主题)"
)

h2("4.3 组件关系图")
para("应用以 main.ts 为入口装配三大基础设施（Pinia、Router、全局守卫），再由路由决定渲染哪一类布局壳，"
     "壳内通过 <router-view> 渲染具体业务页面；业务页面普遍复用通用组件，并通过 Pinia 与全局态交互。")
code_block(
    "App.vue\n"
    " └─ <router-view>\n"
    "     ├─ AdminLayout (后台壳)\n"
    "     │   ├─ AppHeader (用户下拉 / 主题切换 / NoticeBell 通知铃铛)\n"
    "     │   ├─ AppSidebar (菜单树, 由权限 store 驱动)\n"
    "     │   ├─ AppTabs (多页签, 由 app store 驱动)\n"
    "     │   └─ <router-view> + <keep-alive> → 业务页面\n"
    "     │        └─ 复用: DataTable / SearchForm / FormDialog / EChart …\n"
    "     └─ ScreenLayout (大屏壳)\n"
    "         ├─ PlanDashboard (ECharts 大屏)\n"
    "         └─ Simulation3DPage (Three.js 3D 仓库)"
)
screenshot("组件关系图（建议用 draw.io / ProcessOn 按上方层级绘制后插入）")

h2("4.4 数据流图")
para("以「用户登录 → 进入受保护页面 → 拉取列表数据」为例，数据流向清晰单向：")
code_block(
    "① 登录页提交 → userStore.login() → authApi.login() → 后端校验 → 写会话 Cookie\n"
    "② userStore.logged=true 并持久化到 localStorage → 跳转 redirect 目标\n"
    "③ 路由守卫 beforeEach → permissionStore.loadMenu() → 收集权限 Set\n"
    "④ 业务页 onMounted → useRequest(fetcher) → http.get() → 拦截器解包 Result\n"
    "⑤ 数据写入 ref → 模板渲染 → DataTable 分页/操作 → 触发 emit → 重新请求"
)
screenshot("数据流图（登录 → 守卫 → 权限 → 请求 → 渲染 单向数据流）")

h2("4.5 状态管理架构")
para("Pinia 按功能拆分为四个 store，职责单一、边界清晰：")
table(
    ["Store", "职责", "是否持久化"],
    [
        ["user.ts", "登录信息、登录态、login/logout/fetchUserInfo", "是（user、logged）"],
        ["permission.ts", "菜单树、权限 Set、hasPermission getter", "否（Set 不便序列化，登录/刷新后重建）"],
        ["app.ts", "主题、侧栏折叠、多页签", "是（theme、collapsed、tabs）"],
        ["notice.ts", "未读通知数、刷新/重置", "否（实时拉取）"],
    ],
    widths=[1.2, 3.8, 2.0],
)

# ================================================================== 五、核心实现细节
h1("五、核心功能实现细节")

h2("5.1 登录验证与会话鉴权")
para("登录由 userStore.login 统一编排：调用后端 /login（表单编码），成功后置 logged=true 并拉取当前用户信息；"
     "user 与 logged 经持久化插件写入 localStorage，刷新页面登录态不丢失。登出时即便后端接口失败也强制清理本地状态。")
code_block(
    "// stores/user.ts(节选)\n"
    "async login(payload) {\n"
    "  await authApi.login(payload)   // 表单编码 POST /login\n"
    "  this.logged = true\n"
    "  await this.fetchUserInfo()     // GET /admin/user/info\n"
    "}\n"
    "// 持久化:刷新不丢登录态\n"
    "persist: { pick: ['user', 'logged'] }"
)
screenshot("登录页（用户名/密码表单 + 校验提示）")

h2("5.2 路由设计：嵌套 + 动态 + 守卫")
h3("（1）多级嵌套路由")
para("路由按「布局壳 + 子路由」组织：AdminLayout 作为后台父路由，其下挂载 system/user、technology/flow、"
     "plan/order 等 30+ 子路由；ScreenLayout 作为大屏父路由，挂载数据可视化大屏与 3D 仿真。"
     "这构成了 /、/system/user、/digitization/dashboard 等多级路径结构。")

h3("（2）动态路由参数匹配")
para("通知中心实现了真正的路径参数动态路由 /system/notice/:id：收件箱列表「查看」、顶部通知铃铛下拉、"
     "以及 ?open= 深链接都通过 router.push({ name: 'system-notice-detail', params: { id } }) 跳转；"
     "详情页通过 useRoute().params.id 取参拉取详情，并 watch 该参数变化以支持在不同通知间切换时复用组件实例。"
     "此外，路由表末尾以 /:pathMatch(.*)* 通配兜底 404。")
code_block(
    "// router/index.ts(节选)\n"
    "{ path: 'system/notice/:id', name: 'system-notice-detail',\n"
    "  component: () => import('@/views/system/notice/NoticeDetailView.vue'),\n"
    "  meta: { title: '通知详情' } }\n\n"
    "// NoticeDetailView.vue(节选)\n"
    "onMounted(() => load(route.params.id as string))\n"
    "watch(() => route.params.id, (id) => { if (id) load(id as string) })"
)
screenshot("通知详情动态路由页（地址栏可见 /system/notice/<id>）")

h3("（3）路由守卫与权限控制")
para("全局前置守卫 beforeEach 串联三件事：公开页放行、未登录拦截（携带 redirect）、"
     "已登录但权限未加载时重建菜单与权限 Set、最后做路由级 meta.perm 权限校验，"
     "无权则跳 /403。守卫前后由 NProgress 驱动顶部进度条，提供导航反馈。")
code_block(
    "// router/guards.ts(节选)\n"
    "router.beforeEach(async (to) => {\n"
    "  if (to.meta.public) return true\n"
    "  if (!userStore.logged) return { path: '/login', query: { redirect: to.fullPath } }\n"
    "  if (!permStore.loaded) await permStore.loadMenu()\n"
    "  const perm = to.meta.perm\n"
    "  if (perm && !permStore.hasPermission(perm)) return { path: '/403' }\n"
    "  return true\n"
    "})"
)

h2("5.3 Pinia 状态管理与持久化")
para("权限态由 permissionStore 维护：登录或刷新后调用 loadMenu 拉取菜单树，递归收集所有 permission 字符串"
     "（与后端 ShiroRealm 一致按逗号拆分）汇成一个 Set；hasPermission getter 供路由守卫与 <按钮门控> 复用。"
     "Set 不便 JSON 序列化，故权限态不持久化，而是每次进入应用时由菜单树重建，既保证一致性又避免脏数据。")
code_block(
    "// utils/permission.ts(节选)\n"
    "node.permission.split(',').forEach((p) => {\n"
    "  const t = p.trim(); if (t) set.add(t)\n"
    "})"
)
para("应用态由 appStore 维护主题、侧栏折叠、多页签并全部持久化；主题在 main.ts 启动时调用 applyTheme 应用到 "
     "<html>，实现刷新后明暗主题保持。")

h2("5.4 组合式 API 与生命周期钩子")
para("项目全面采用 <script setup> 组合式 API。生命周期钩子在 14+ 个组件中使用：onMounted 用于页面首次加载数据、"
     "onUnmounted/onBeforeUnmount 用于销毁 ECharts 实例、Three.js 渲染循环与 SSE 连接，onActivated 配合 keep-alive。"
     "逻辑复用以 composable 形式沉淀：")
bullet("useRequest：统一管理异步请求的 loading / error / data，配合骨架屏与错误反馈。")
bullet("usePagination：封装 current / size / total 分页状态与重置。")
bullet("useAiChat：封装 AI 助手的消息流、发送、中断、历史拼接。")
bullet("useDict / useTypewriter / useFloatingWindow：字典下拉、打字机效果、悬浮窗拖拽。")
code_block(
    "// composables/useRequest.ts(节选)\n"
    "async function run(...args) {\n"
    "  loading.value = true; error.value = null\n"
    "  try { data.value = await fetcher(...args); options.onSuccess?.(data.value) }\n"
    "  catch (e) { error.value = e; throw e }\n"
    "  finally { loading.value = false }\n"
    "}"
)

h2("5.5 组件化：通用组件与组件通信")
para("通用组件与业务组件分层，最大化复用。通用组件以 props 入、$emit 出，业务页面仅负责装配与数据请求：")
table(
    ["通用组件", "职责", "通信方式"],
    [
        ["DataTable", "服务端分页表格 + 工具栏插槽 + 操作列插槽", "props: data/columns/pager；emit: page-change/size-change"],
        ["SearchForm", "内联查询表单 + 搜索/重置按钮", "props: model；emit: search/reset；默认插槽放表单项"],
        ["FormDialog", "弹窗表单壳（替代 iframe 弹层）", "props: visible/title；emit: update/confirm"],
        ["DualListTransfer / OrderedTransfer", "穿梭框（工艺路线选工序、排序）", "props: 源/目标；emit: change"],
        ["EChart", "ECharts 通用封装（自适应 + 销毁）", "props: option；内部 onUnmounted dispose"],
        ["ImageUpload / MultiImageUpload", "物料图片上传", "props: modelValue；emit: update:modelValue"],
        ["TreeTable / MasterDetailLayout", "树形表格 / 主从布局", "props 驱动 + 默认插槽"],
        ["NoticeBell", "通知铃铛（未读角标 + 下拉）", "读取 notice store + router 跳转动态路由"],
    ],
    widths=[1.9, 2.6, 2.9],
)
para("组件间通信遵循三条路径：父子用 props / $emit（如 DataTable 翻页回调）、"
     "跨层级全局态用 Pinia（如登录态、权限、未读数、主题）、URL 状态用路由参数（如通知详情 :id）。"
     "数据流向单一、可追踪。")
screenshot("某 CRUD 业务页（SearchForm + DataTable + FormDialog 组合）")

h2("5.6 Axios 封装与请求拦截")
para("api/request.ts 封装统一的 axios 实例与拦截器，屏蔽后端约定细节：")
bullet("请求拦截：POST 默认把 JSON 转 application/x-www-form-urlencoded（适配后端表单接口）；"
       "显式声明 application/json 或二进制（FormData/Blob）则跳过。")
bullet("响应拦截：解包后端统一 Result 包裹——code===0 返回 data，否则弹出 ElMessage 错误并抛出；"
       "HTTP 401 自动跳转登录页。")
bullet("会话：withCredentials 携带 Shiro 会话 Cookie；X-Requested-With 头确保 401 返回 JSON。")
code_block(
    "// api/request.ts(节选)\n"
    "service.interceptors.response.use(\n"
    "  (resp) => { try { return unwrapResult(resp.data) }\n"
    "    catch (e) { ElMessage.error(e.message); return Promise.reject(e) } },\n"
    "  (error) => { if (error.response?.status === 401) location.href = '/login'\n"
    "    return Promise.reject(error) }\n"
    ")"
)

h2("5.7 交互体验：加载状态与表单校验")
bullet("加载状态：列表用骨架屏（TableSkeleton / el-skeleton）与表格 loading；路由切换用 NProgress 顶部进度条；"
       "AI 助手用打字机逐字渲染。")
bullet("表单校验：Element Plus el-form 的 rules 提供必填、长度、正则校验与即时错误反馈；"
       "提交前 validate 拦截非法输入，避免脏数据落库。")
bullet("反馈：所有写操作经 ElMessage / ElMessageBox 给出成功提示或二次确认（如删除确认）。")
screenshot("骨架屏加载态 与 表单校验错误提示")

# ================================================================== 六、创新与深度优化
h1("六、创新与深度优化")

h2("6.1 数据可视化大屏（ECharts）")
para("digitization/dashboard 是独立于后台壳的全屏看板（ScreenLayout），用 ECharts + vue-echarts 渲染"
     "产能、订单、库存等多维统计图表，配合定时轮询近实时刷新。EChart 通用组件统一处理自适应 resize 与"
     "组件卸载时的实例 dispose，避免内存泄漏。")
screenshot("数据可视化大屏（多图表全屏看板）")

h2("6.2 3D 数字孪生仓库（Three.js）")
para("digitization/simulation 用 Three.js 构建三维数字孪生仓库：按库位坐标生成货架与库存方块，"
     "依据库存量映射热力颜色（heatColor 工具），支持鼠标点选库位查看明细。渲染循环在组件卸载时被正确停止，"
     "WebGL 上下文按需释放。")
screenshot("3D 数字孪生仓库（三维货架 + 热力着色 + 点选）")

h2("6.3 AI 智能助手（SSE 流式）")
para("AI 助手以悬浮球 + 悬浮窗形式常驻（AiFab / AiChatWindow），对接后端 Agent 端点 /admin/ai/chat。"
     "由于流式 SSE 需要手动分帧并处理 401，前端绕开 axios 直接用 fetch + ReadableStream 读取，"
     "经 createSseParser 解析事件，applyAiEvent 归约到消息状态，并以打字机效果逐字呈现；"
     "AbortController 支持随时中断生成。")
code_block(
    "// api/ai.ts(节选)\n"
    "resp = await fetch(`${base}/admin/ai/chat`, {\n"
    "  method: 'POST', headers: { 'Content-Type': 'application/json' },\n"
    "  body: JSON.stringify({ messages }), credentials: 'include', signal })\n"
    "const reader = resp.body.getReader()  // 手动分帧解析 SSE"
)
screenshot("AI 助手悬浮窗（流式回答 + 工具调用步骤）")

h2("6.4 BPMN 流程设计器与轻量审批引擎")
para("工作流模块用 bpmn-js 集成可视化 BPMN 设计器（diagram-js-minimap 缩略图），"
     "可拖拽绘制流程模型；后端配套轻量 BPMN 解析与运行时（start/claim/complete/reject），"
     "生产订单提交后派生审批待办，构成「录入 → 审批 → 运算」的闭环。")
screenshot("BPMN 流程模型设计器（全屏画布 + 小地图）")

h2("6.5 生产甘特图与通知中心")
bullet("生产甘特图：order/gantt 自研甘特排程，支持拖拽调整开工/完工、进度展示。")
bullet("通知中心：发布端（全员/角色/指定用户 + 已读统计）+ 接收端（收件箱筛选/详情/已读/删除）+ "
       "顶部铃铛未读角标，详情页采用动态路由 /system/notice/:id。")
screenshot("生产甘特图 + 通知中心收件箱")

h2("6.6 性能优化")
table(
    ["优化手段", "实现方式", "目标"],
    [
        ["路由懒加载", "所有路由 component 用 () => import() 动态导入", "按页加载，减小首屏体积"],
        ["组件按需引入", "unplugin-auto-import + unplugin-vue-components 自动 tree-shaking", "只打包用到的 Element/Vue API"],
        ["产物分包", "Vite manualChunks 将 element / vue 框架拆为独立 chunk", "利于浏览器长期缓存"],
        ["组件缓存", "AdminLayout 内 <keep-alive :max=12> 缓存已访问页面", "页签切换免重渲染、状态保留"],
        ["顶部进度条", "NProgress 在守卫前后驱动", "导航期视觉反馈"],
        ["重型库懒加载", "ECharts / Three.js / bpmn-js 仅在对应路由按需加载", "首屏不被大库拖慢"],
    ],
    widths=[1.6, 3.6, 1.8],
)
para("性能分析使用 Chrome DevTools Performance 面板与 Lighthouse：定位长任务与首屏瓶颈，"
     "据此实施上述分包与懒加载策略，首屏加载时间控制在 3 秒以内。")
screenshot("Lighthouse / DevTools 性能评分截图")

# ================================================================== 七、工程化
h1("七、代码质量与工程化")

h2("7.1 目录规范与代码风格")
bullet("遵循 Vite 模块化目录（api / components / composables / stores / router / views / utils / types）。")
bullet("ESLint 10 + Prettier 统一缩进、引号、分号；npm run lint 一键修复；构建期 vue-tsc 做类型检查。")
bullet("命名语义化：组件 PascalCase、composable 以 use 前缀、工具函数动词命名；关键逻辑均有中文注释。")

h2("7.2 版本控制与分支策略")
para("代码托管于 Git（含 Gitee 远程），采用 main / develop / feature 三层分支策略："
     "main 为稳定线，develop 为集成线，各功能在 feature/* 分支开发完成后合并回 develop。")
bullet("提交规范：采用 emoji + conventional commit（如 ✨ feat、🐛 fix、📝 docs、♻️ refactor、✅ test），"
       "提交信息使用中文，语义清晰。")
bullet("提交节奏：按模块/页面规律增量提交，仓库累计 800+ 次提交，无一次性巨型提交；每周提交远超 2 次。")
screenshot("Git 提交历史 / 分支图（gitk 或 IDE Git 面板）")

h2("7.3 单元测试")
para("使用 Vitest 4 对纯逻辑与组合式函数做单元测试，覆盖 SSE 分帧、AI 事件归约、打字机、权限收集、"
     "BPMN 解析、甘特计算、各业务工具函数等关键模块。当前 30 个测试文件、333 个用例全部通过。")
code_block("$ npx vitest run\n Test Files  30 passed (30)\n      Tests  333 passed (333)")

# ================================================================== 八、评分对照
h1("八、评分细则对照")

para("下表把作业五个评分维度逐条映射到本项目的落地实现，便于核对：")
table(
    ["评分维度（20 分/项）", "本项目对应实现"],
    [
        ["① 功能交互与实现",
         "30+ 业务页面全 SFC 开发；通用组件(表单/表格/弹窗/穿梭框)与业务组件分层；"
         "props/$emit + Pinia 通信；骨架屏 + Loading + NProgress；Element Plus 正则校验与错误反馈；"
         "unplugin 全量按需引入；关键交互 < 1s"],
        ["② 代码质量与工程化",
         "Vite 模块化目录；ESLint + Prettier 统一风格；语义化命名 + 中文注释；"
         "main/develop/feature 分支；emoji conventional commit；800+ 次规律提交"],
        ["③ 路由与状态管理",
         "多级嵌套路由(布局壳 + 子路由)；动态路由 /system/notice/:id；beforeEach 登录拦截 + 权限校验；"
         "Pinia 模块化(user/permission/app/notice)；localStorage 持久化刷新不丢登录态与主题"],
        ["④ 创新与深度优化",
         "5 大亮点：ECharts 大屏 / Three.js 3D 数字孪生 / AI 助手 SSE / BPMN 设计器 / 甘特图；"
         "DevTools + Lighthouse 性能分析；懒加载 + 按需 + 分包 + keep-alive；首屏 < 3s"],
        ["⑤ 项目文档",
         "本文档：需求分析 + 技术选型 + 系统架构(组件关系图/数据流图) + 关键功能技术方案 + 评分对照，"
         "图文并茂、5000 字以上"],
    ],
    widths=[1.9, 5.1],
)

# ================================================================== 九、部署运行
h1("九、部署与运行")
h3("开发环境")
code_block(
    "cd mes/vue3\n"
    "pnpm install          # 安装依赖\n"
    "pnpm dev              # 启动开发服务器(端口 4200, /api 代理至 :9090)\n"
    "pnpm typecheck        # vue-tsc 类型检查\n"
    "pnpm test             # 运行 Vitest 单元测试\n"
    "pnpm lint             # ESLint 代码检查并修复"
)
h3("生产构建")
code_block(
    "pnpm build            # vue-tsc 构建 + Vite 打包(产物在 dist/)\n"
    "pnpm preview          # 本地预览生产产物\n"
    "# 生产环境由 Nginx 反向代理:静态资源 + /api 转发至 Spring Boot"
)
para("默认演示账号：admin / 123（开发环境已关闭验证码，便于脚本化登录与演示）。")
screenshot("本地启动后的运行效果（终端 + 浏览器）")

# ================================================================== 十、总结
h1("十、总结与展望")
para("本项目以 Vue3 组合式 API 为核心，完整实现了一套面向工业互联网的制造执行系统前端，"
     "覆盖系统管理、基础数据、工艺、计划、库存、数字化、工作流八大业务域共 30+ 业务页面。"
     "工程上严格落地了作业要求的全部硬性指标：登录验证、组合式 API 生命周期钩子、Pinia、Vue Router 4、"
     "Element Plus；并在此之上扩展了 ECharts 大屏、Three.js 数字孪生、AI 流式助手、BPMN 设计器、"
     "生产甘特图等多个创新亮点。")
para("代码质量方面，项目遵循模块化目录与统一代码风格，配套 333 个单元测试用例与规范的 Git 分支/提交策略，"
     "具备良好的可维护性与协作性。性能方面，通过路由懒加载、组件按需引入、产物分包与 keep-alive 缓存，"
     "将首屏与关键交互控制在目标区间内。")
para("后续可进一步拓展：引入 WebSocket 替代轮询实现真正实时推送、补充 E2E 测试、接入 CDN 加速静态资源、"
     "以及把 AI 助手的工具调用能力扩展到更多业务场景。")

# ------------------------------------------------------------------ 保存
doc.save(OUTPUT)
print(f"已生成：{OUTPUT}")
